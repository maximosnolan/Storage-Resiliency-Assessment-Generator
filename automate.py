from docx import Document
from numpy import NaN
import pandas
import time
from docx.shared import Cm, Pt, RGBColor
from copy import deepcopy
from pandas.io import excel
from multiprocessing import Pool
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Recomendation:
    def __init__(self, importance, recommendation, framework, rank, cat, cat_rank):
        self.importance = importance #Importance (range from 1-5)
        self.recommendation = recommendation #recommendation (Text of the recommendation)
        self.framework = framework #framework (from NIST)
        self.rank = rank #rank (used from 1-5, in order of NIST framework, used for sorting recommendations)
        self.cat = cat #cat (category this falls into, used only for the executive summary % achieved graph)
        self.cat_rank = cat_rank #cat_rank (used in ordering recommendations for executive summary)


class company:
    def __init__(self, assessment, name, date):
        self.assessment = assessment #the file of the assessment (XLSX file)
        self.company_name = name #the name of the company
        self.date = date #date entered
    date = ""
    company_name = ""
    assessment = ""


major_reccomendation = [] #major recomendations
minor_reccomendation = [] #minor recomendations
final_scores = {} #scores of % achieved, hash map which maps the name of the metric to the score achieved
num_rows = 100

def get_information():
    df = pandas.read_excel('CRAT_example.xlsx', sheet_name= 'Customer Basic Information')
    #assessment = input("Please enter the assessment file name:")
    #name = input("Enter the companies name:")
    #date = input("Enter the date the assessment will be presented")
    assessment = 'CRAT_example.xlsx'
    name = df.loc[9, 'Information']
    date = df.loc[4, 'Information']
    victim_comapny = company(assessment, name, date)
    return victim_comapny

def print_intro():
    print("Welcome to the IBM Storage Resiliency Report Automation Software!")
    print("Use guide:")
    print("Please have the assessment in the folder where you run this program and the templated word doc provided (Is called 'assessment_template.docx').")
    print("As well, you will also need screenshots of the images. Don't worry about their size, the script will fit them for the document, name them as:")
    image_map = {1: "Maturity_Graph_1", 2: "Maturity_Graph_2", 3: "Ransom", 4: "Executive"}
    for x, y in image_map.items():
        print("Image #" , x , " as: ", y)
    print("Once you run the script a new document will appear named 'IBM Storage Assessment (company_name)'")
    return

def replace_all(company, document, randsom_percent, final_scores):
    document.add_paragraph(company.company_name)
    #settig style for the paragraph
    style = document.styles['Normal']
    font = style.font
    font.name = 'IBM Plex Sans Light'
    font.size = Pt(9)
    for paragraph in document.paragraphs:
        if 'Executive Summary – Summary View' in paragraph.text:
            df = copy_total_score()
            document.tables[1].cell(1,1).text = str(round(df.loc[2, 'SCORE'], 2))
            for x in range (4,35):
                #pulling information for executive summary, due to the way it is formatted in the xlsx file, we need to skip a few lines
                if x != 10 and x != 17 and x != 22 and x != 29 and x!= 35:
                   document.tables[1].cell(x - 1, 1).text = str(round(df.loc[x, 'SCORE'], 2))
                   document.tables[1].cell(x - 1, 1).allignment = 1
                   document.tables[1].cell(x - 1, 2).text = str(df.loc[x, 'LEVEL'])
                   document.tables[1].cell(x - 1, 2).allignment = 1
        if 'CUSTOMER_NAME' in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "CUSTOMER_NAME", company.company_name)
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'PRESENTATION_DATE' in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "PRESENTATION_DATE", company.date)
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'PERCENT_RANSOM' in paragraph.text:
            orig_text = paragraph.text
            percent = str(round(randsom_percent,2 )) + '%'
            new_text = str.replace(orig_text, "PERCENT_RANSOM", percent)
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'PERCENT_MISS' in paragraph.text:
            orig_text = paragraph.text
            miss = round(100 - int(randsom_percent),2)
            missing = str(miss) + '%'
            new_text = str.replace(orig_text, "PERCENT_MISS", missing)
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'EX_1' in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "EX_1", list(final_scores.keys())[15])
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'EX_2' in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "EX_2", list(final_scores.keys())[16])
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'EX_3' in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "EX_3", list(final_scores.keys())[17])
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
        if 'SCORE  is slightly below average' in paragraph.text:
            orig_text = paragraph.text
            conjunction = ""
            if round(df.loc[2, 'SCORE'], 2) < 6:
                conjunction = " is below average"
            elif round(df.loc[2, 'SCORE'], 2) < 7:
                conjunction = " is slightly below average"
            elif round(df.loc[2, 'SCORE'], 2) > 9:
                conjunction = " is above average"
            elif round(df.loc[2, 'SCORE'], 2) > 8:
                conjunction = " is slightly above average"
            else:
                conjunction = " is average"
            new_text = str.replace(orig_text, "SCORE  is slightly below average",str(round(df.loc[2, 'SCORE'], 2)) + conjunction )
            paragraph.text = new_text
            style = document.styles['Normal']
            font = style.font
            font.name = 'IBM Plex Sans Light'
            font.size = Pt(9)
            paragraph.style = document.styles['Normal']
    x = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if x == 0: #if this is the first table
                        continue
                    if 'CUSTOMER_NAME' in paragraph.text:
                        paragraph.text = paragraph.text.replace("CUSTOMER_NAME", company.company_name)
                    if 'PRESENTATION_DATE' in paragraph.text:
                        paragraph.text = paragraph.text.replace("PRESENTATION_DATE", company.date)
        x+=1

    
    return

def copy_total_score():
    return pandas.read_excel('CRAT_example.xlsx', sheet_name= 'Results')
def pull_level(score):
    score_map = {"Identify": 1, "Protect": 2, "Detect": 3, "Respond": 4, "Recover": 5 }
    return score_map[score]
def pull_framework(level):
    #set framework depending on depth in assessment
    if level < 50: return "Identify", 1
    elif level <109: return "Protect", 2
    elif level < 146: return "Detect", 3
    elif level < 192: return "Respond", 4
    else: return "Recover", 5

def pull_reccomendations():
    excel_data_df = pandas.read_excel('CRAT_example.xlsx', sheet_name= 'Internal Data')
    #227 corresponds to first question "Inventory storage physical devices...."
    for x in range (227, 399):
        #1 determine if yes 
        level = pull_level(str(excel_data_df.loc[x, 'Score']))
        if str(excel_data_df.loc[x, 'Answer']) == 'Yes':
            continue
        elif str(excel_data_df.loc[x, 'Answer']) == 'No':
            a = Recomendation(str(excel_data_df.loc[x, 'Count']), str(excel_data_df.loc[x, 'Questions']), str(excel_data_df.loc[x, 'Score']), level, "0", 0)
            major_reccomendation.append(a)
        else:
            a = Recomendation(str(excel_data_df.loc[x, 'Count']), str(excel_data_df.loc[x, 'Questions']), str(excel_data_df.loc[x, 'Score']), level, "0", 0)
            minor_reccomendation.append(a)

    return

def pull_score():
    scores = {}
    excel_data_df = pandas.read_excel('CRAT_example.xlsx', sheet_name= 'Internal Data')
    for x in range (1, 19):
        scores[excel_data_df.loc[x, 'Category']] = excel_data_df.loc[x, 'Percent_Ach']
    final_scores = {k: v for k, v in sorted(scores.items(), key=lambda item: item[1])}
    return final_scores

def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)

def generate_table(document, maj_min, number):
    table = document.tables[number]
    for x in range(0, len(maj_min) ):
        if x > num_rows: #if there are more than 60 reccomendations
           print("More than:", num_rows, "recommendations, please add more rows to the template")
           continue
        rec = maj_min[x]
        importance = rec.importance #int from 1-5
        text = rec.recommendation #text
        type = rec.framework #NIST
        document.tables[number].cell(x,0).text = importance
        document.tables[number].cell(x,0).paragraphs[0].runs[0].font.size = Pt(9)
        document.tables[number].cell(x,0).paragraphs[0].runs[0].font.name = 'Arial'
        
        document.tables[number].cell(x,1).text = text
        document.tables[number].cell(x,1).paragraphs[0].runs[0].font.size = Pt(9)
        document.tables[number].cell(x,1).paragraphs[0].runs[0].font.name = 'Arial'
        
        
        if type == "Identify":
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w')))
            document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w'))))
            p2 = table.cell(x,1)
            p2._tc.get_or_add_tcPr().append(shading_elm)
        elif type == "Protect":
            shading_elm = parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w')))
            p1 = table.cell(x,0)
            p2 = table.cell(x,1)
            document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w'))))
            p2._tc.get_or_add_tcPr().append(shading_elm)
        elif type == "Respond":
            shading_elm = parse_xml(r'<w:shd {} w:fill="E6B9B7"/>'.format(nsdecls('w')))
            p1 = table.cell(x,0)
            p2 = table.cell(x,1)
            document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E6B9B7"/>'.format(nsdecls('w'))))
            p2._tc.get_or_add_tcPr().append(shading_elm)
        elif type == "Detect":
            shading_elm = parse_xml(r'<w:shd {} w:fill="FCD6B4"/>'.format(nsdecls('w')))
            p1 = table.cell(x,0)
            p2 = table.cell(x,1)
            document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FCD6B4"/>'.format(nsdecls('w'))))
            p2._tc.get_or_add_tcPr().append(shading_elm)
        else: #is recover
            shading_elm = parse_xml(r'<w:shd {} w:fill="ECF2DF"/>'.format(nsdecls('w')))
            p1 = table.cell(x,0)
            p2 = table.cell(x,1)
            document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="ECF2DF"/>'.format(nsdecls('w'))))
            p2._tc.get_or_add_tcPr().append(shading_elm)
    #remove extra rows
    for x in range(len(maj_min), num_rows): #Change if you add more rows
        row = table.rows[len(maj_min)]
        remove_row(table, row)
    return

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)
    return

def sort_reccomendations(maj_min):
    maj_min.sort(key = lambda x: x.importance)
    maj_min.sort(key = lambda x: x.rank)
    return

def set_font_all(document, paragraph_num):
    document.paragraph[paragraph_num].font.name = 'IBM Plex Sans Light'
    document.paragraph[paragraph_num].font.size = Pt(9)
    return

def save_images(document, number, pic):
    table = document.tables[number]
    row_cells = table.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run()
    if pic == "Ransom.jpg" or pic == "Executive.jpg":
        table.rows[1].height = Cm(11.5)
        run.add_picture(pic, width = 5400000, height = 4000000)
    if pic == "Maturity_Graph_1.jpg" or pic == "Maturity_Graph_2.jpg":
        if pic == "Maturity_Graph_1.jpg":
            table.rows[1].height = Cm(10.39)
            run.add_picture(pic, width = 5400000, height = 3500000)
        else: 
            table.rows[1].height = Cm(12)
            run.add_picture(pic, width = 5400000, height = 3900000)
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col in table.columns:
        for cell in col.cells:
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return

def add_footer(document, company):
    for section in document.sections:
        footer_section = section
        footer = footer_section.footer
        footer_text = footer.paragraphs[0]
        footer_text.text = "\tIBM -" + company.company_name + " CONFIDENTIAL"
        footer.paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
        footer.paragraphs[0].runs[0].font.size = Pt(9)
   

    return

def pull_ransom():
    excel_data_df = pandas.read_excel('CRAT_example.xlsx', sheet_name= 'Internal Data')
    return round(((excel_data_df.loc[6, 'Ransom_percent'])  * 100), 2)
   
def change_intro(document, company):
    #access table 0
    count = 0
    table = document.tables[0]
    
    document.tables[0].cell(1,0).text = '        ' + company.company_name
    document.tables[0].cell(1,0).paragraphs[0].runs[0].font.size = Pt(23)
    document.tables[0].cell(1,0).paragraphs[0].runs[0].font.bold = True
    document.tables[0].cell(1,0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    document.tables[0].cell(1,0).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'

    
    document.tables[0].cell(3,0).text = '                        '  + company.date
    document.tables[0].cell(3,0).paragraphs[0].runs[0].font.size = Pt(14)
    document.tables[0].cell(3,0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    document.tables[0].cell(3,0).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
    
    return

def insights(document, number, insight_type, final_scores, num_to_pull):
    #set number equal to what table we are setting. In the case that more tables are added/removed, edit number to edit a different table
    #DEBUG
    major_rec_insights = []
    minor_rec_insights = [] 
    excel_data_df = pandas.read_excel('CRAT_example.xlsx', sheet_name= 'Internal Data')
    if(insight_type == "Ransom"):
        for x in range (2, 225):
            if excel_data_df.loc[x, 'Ransomware?'] == 'x': #if this question relates to ransomware
                if excel_data_df.loc[x, 'Answer'] == 'Yes': #if yes, we ignore
                    continue
                elif excel_data_df.loc[x, 'Answer'] == 'No':
                    #framework is set as color
                    init_identity, init_rank = pull_framework(x)
                    victim_rec = Recomendation(str(excel_data_df.loc[x, 'Rank']), str(excel_data_df.loc[x, 'Questions']), init_identity,init_rank, "0", 0)
                    major_rec_insights.append(victim_rec)
                else: #is high/partial/low partail answer, minor reccomendation
                    init_identity, init_rank = pull_framework(x)
                    victim_rec = Recomendation(str(excel_data_df.loc[x, 'Rank']), str(excel_data_df.loc[x, 'Questions']), init_identity, init_rank, "0", 0)
                    minor_rec_insights.append(victim_rec)
    if(insight_type == "Executive"):
          for x in range (2, 225):
            category = ""
            if executive_help(excel_data_df, x, final_scores, "Category 1", num_to_pull) or executive_help(excel_data_df, x, final_scores, "Category 2", num_to_pull): #if this question relates to the lowest four
                trash, cat_rank, cat = executive_helper(excel_data_df, x, final_scores, "Category 1", num_to_pull)
                if executive_help(excel_data_df, x, final_scores, "Category 1", num_to_pull) and executive_help(excel_data_df, x, final_scores, "Category 2", num_to_pull):
                    category = str(excel_data_df.loc[x, "Category 1"]) + ", " + str(excel_data_df.loc[x, "Category 2"])
                if trash == False:
                    trash, cat_rank, cat = executive_helper(excel_data_df, x, final_scores, "Category 2", num_to_pull)
                if excel_data_df.loc[x, 'Answer'] == 'Yes': #if yes, we ignore
                    continue
                elif excel_data_df.loc[x, 'Answer'] == 'No':
                    init_identity, init_rank = pull_framework(x)
                    if(category == ""):
                        category = str(excel_data_df.loc[x, cat])
                    victim_rec = Recomendation(      str(int(excel_data_df.loc[x, 'Rank'])), str(excel_data_df.loc[x, 'Questions']), init_identity,init_rank, category, cat_rank)
                    major_rec_insights.append(victim_rec)
                else: #is high/partial/low partail answer, minor reccomendation
                    init_identity, init_rank = pull_framework(x)
                    if(category == ""):
                        category = str(excel_data_df.loc[x, cat])
                    temp = Recomendation(str(int(excel_data_df.loc[x, 'Rank'])), str(excel_data_df.loc[x, 'Questions']), init_identity, init_rank, category, cat_rank)
                    minor_rec_insights.append(temp)
    #reccomendations pulled, add to table!
    table = document.tables[number]
    x = 0 #this denotes where we start inputing for minor reccomendations
    recommendation_list = major_rec_insights + minor_rec_insights
    sort_reccomendations(recommendation_list)
    if(insight_type == "Executive"): recommendation_list.sort(key = lambda x: x.cat_rank)
    for reccomendation in recommendation_list:
            if x > num_rows: #if there are more than 60 reccomendations
                print("More than:", num_rows, "recommendations, please add more rows to the template")
                continue
            #Set Color
            significance = "Major"
            if reccomendation in major_rec_insights:
                significance = "Major"
            else: significance = "Minor"
            type = reccomendation.framework
            
            table.cell(x,0).text = reccomendation.importance[0]
            table.cell(x,0).paragraphs[0].runs[0].font.size = Pt(9)
            table.cell(x,0).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
            table.cell(x,1).text = reccomendation.recommendation
            table.cell(x,1).paragraphs[0].runs[0].font.size = Pt(9)
            table.cell(x,1).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
            table.cell(x,2).text = significance
            table.cell(x,2).paragraphs[0].runs[0].font.size = Pt(9)
            table.cell(x,2).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
            num_columns = 2 if insight_type == "Ransom" else 3
            if insight_type == "Executive": #if executive, need to set metric column
                table.cell(x,3).text = reccomendation.cat
                table.cell(x,3).paragraphs[0].runs[0].font.size = Pt(9)
                table.cell(x,3).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
            if type == "Identify":
                table.cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w'))))
                table.cell(x,1)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w'))))
                table.cell(x,2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w'))))
                if insight_type == "Executive": document.tables[number].cell(x,3)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w'))))
            elif type == "Protect":
                document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w'))))
                document.tables[number].cell(x,1)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w'))))
                document.tables[number].cell(x,2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w'))))
                if insight_type == "Executive": document.tables[number].cell(x,3)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w'))))
                #set_cell_color(table, x, num_columns, parse_xml(r'<w:shd {} w:fill="E4DFEC"/>'.format(nsdecls('w'))))
            elif type == "Respond":
                shading_elm = parse_xml(r'<w:shd {} w:fill="E6B9B7"/>'.format(nsdecls('w')))
                p1 = table.cell(x,0)
                p2 = table.cell(x,1)
                document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E6B9B7"/>'.format(nsdecls('w'))))
                document.tables[number].cell(x,2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E6B9B7"/>'.format(nsdecls('w'))))
                p2._tc.get_or_add_tcPr().append(shading_elm)
                if insight_type == "Executive": 
                    document.tables[number].cell(x,3)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E6B9B7"/>'.format(nsdecls('w'))))
            elif type == "Detect":
                shading_elm = parse_xml(r'<w:shd {} w:fill="FCD6B4"/>'.format(nsdecls('w')))
                p1 = table.cell(x,0)
                p2 = table.cell(x,1)
                document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FCD6B4"/>'.format(nsdecls('w'))))
                document.tables[number].cell(x,2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FCD6B4"/>'.format(nsdecls('w'))))
                p2._tc.get_or_add_tcPr().append(shading_elm)
                if insight_type == "Executive": 
                    
                    document.tables[number].cell(x,3)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FCD6B4"/>'.format(nsdecls('w'))))
            else: #is recover
                shading_elm = parse_xml(r'<w:shd {} w:fill="ECF2DF"/>'.format(nsdecls('w')))
                p1 = table.cell(x,0)
                p2 = table.cell(x,1)
                document.tables[number].cell(x,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="ECF2DF"/>'.format(nsdecls('w'))))
                document.tables[number].cell(x,2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="ECF2DF"/>'.format(nsdecls('w'))))
                p2._tc.get_or_add_tcPr().append(shading_elm)
                if insight_type == "Executive": 
                   
                    document.tables[number].cell(x,3)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="ECF2DF"/>'.format(nsdecls('w'))))
            x+=1 #Incriment line count

    #tables should be set now. 
    for y in range(x, num_rows): #Change if you add more rows
        row = table.rows[x]
        remove_row(table, row)
    #fix header table sizes
    for x in range(3):
        document.tables[5].cell(0,x).paragraphs[0].runs[0].font.size = Pt(12)
        document.tables[5].cell(0,x).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
    for x in range(4):
        document.tables[8].cell(0,x).paragraphs[0].runs[0].font.size = Pt(12)
        document.tables[8].cell(0,x).paragraphs[0].runs[0].font.name = 'IBM Plex Sans Light'
    return

def executive_help(df, x, final_scores, cat, num_elements):
    for count in range(num_elements):
        if df.loc[x, cat] == list(final_scores.keys())[count]:
            return True,
    return False

def executive_helper(df, x, final_scores, cat, num_elements):
    for count in range(num_elements):
        if df.loc[x, cat] == list(final_scores.keys())[count]:
            return True, count, cat
    return False, 0, "nocat"

def set_cell_color(table, row_number, num_columns, color_lambda):
    #loops don't work for this, need to hard code in different cases
    shading_element = color_lambda
    if(num_columns >= 3):
        table.cell(row_number, 1)._tc.get_or_add_tcPr().append(color_lambda)
    if num_columns >= 4:
        table.cell(row_number, 2)._tc.get_or_add_tcPr().append(color_lambda)
    if num_columns >= 5:
        table.cell(row_number, 3)._tc.get_or_add_tcPr().append(color_lambda)
    p1 = table.cell(row_number, 1)
    p2 = table.cell(row_number, 2)
    p1._tc.get_or_add_tcPr().append(shading_element)
    p2._tc.get_or_add_tcPr().append(shading_element)
    return



def save_image_driver(document):
    #save images to word document
    save_images(document, 4, "Ransom.jpg")
    save_images(document, 2, "Maturity_Graph_1.jpg")
    save_images(document, 3, "Maturity_Graph_2.jpg")
    save_images(document, 7, "Executive.jpg")
    return

"""def pull_image():
    app = Dispatch("Excel.Application")
    workbook_file_name = 'CRAT_example.xlsx'
    workbook = app.Workbooks.Open(Filename=workbook_file_name)
    # WARNING: The following line will cause the script to discard any unsaved changes in your workbook
    app.DisplayAlerts = False
    i = 1
    for sheet in workbook.Worksheets:
        for chartObject in sheet.ChartObjects():
            print(sheet.Name + ':' + chartObject.Name)
            chartObject.Chart.Export("chart" + str(i) + ".png")
            i += 1
    workbook.Close(SaveChanges=False, Filename=workbook_file_name)"""

def main():
    print_intro()
    victim_company = get_information()
    start_time = time.time()
    document = Document('assessment_template.docx')
    print("Processing Major recommendations...")
    pull_reccomendations()
    sort_reccomendations(major_reccomendation)
    sort_reccomendations(minor_reccomendation)
    print(len(minor_reccomendation))
    final_scores = pull_score()
    print("Processing Minor recommendations...")
    generate_table(document, major_reccomendation, 11)
    generate_table(document, minor_reccomendation, 13)
    ransom_percent = pull_ransom()
    print("Auto populating company spisific information...")
    replace_all(victim_company, document, ransom_percent, final_scores)
    copy_total_score()
    print("Importing Images and fitting them...")
    save_image_driver(document= document) 
    add_footer(document, victim_company)
    print("Formating text and tables...")
    change_intro(document, victim_company)
    num_to_pull = 4 #set to determine how many catagories are in executive summary reccoemdations
    print("adding insights...")
    insights(document= document, number = 6,insight_type=  "Ransom",final_scores= final_scores, num_to_pull= num_to_pull)
    insights(document= document, number = 9,insight_type= "Executive", final_scores =final_scores, num_to_pull= num_to_pull)
    #pull_image() WIP, will revisit later
    print("Compilation process completed in: ",  time.time() - start_time, " seconds.")
    document_name =  victim_company.company_name + "CRAT_Assessment--Final.docx"
    print(document_name)
    print("New document saved as:", document_name)
    document.save(document_name)
    #open = 'start ' + document_name
    #os.system(open)
    #F = open(document_name, "w")
    return

if __name__ == "__main__":
    main()
